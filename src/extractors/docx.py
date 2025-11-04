"""DOCX file extractor."""

import logging
from pathlib import Path

from docx import Document

from src.extractors.base import DocumentExtractor, ExtractedDocument

logger = logging.getLogger(__name__)


class DocxExtractor(DocumentExtractor):
    """Extractor for Microsoft Word DOCX files."""

    def extract(self, file_path: str) -> ExtractedDocument:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            ExtractedDocument with extracted text
        """
        path = Path(file_path)

        try:
            doc = Document(str(path))

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)

            if table_texts:
                content += "\n\n" + "\n".join(table_texts)

            # Build metadata
            metadata = self.get_file_info(file_path)
            metadata["extractor"] = "docx"
            metadata["num_paragraphs"] = len(doc.paragraphs)
            metadata["num_tables"] = len(doc.tables)

            # Add document properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject

            return ExtractedDocument(
                content=content, metadata=metadata, source=str(path.absolute())
            )
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise

    def supports(self, file_path: str) -> bool:
        """Check if file is a DOCX.

        Args:
            file_path: Path to file

        Returns:
            True if file is DOCX
        """
        return Path(file_path).suffix.lower() == ".docx"

